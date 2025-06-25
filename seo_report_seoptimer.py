import streamlit as st
import requests
import json
import io
from docx import Document
from dotenv import load_dotenv
import os

# ───── Load environment variables ─────
load_dotenv()
API_KEY = os.getenv("PAGESPEED_API_KEY")

if not API_KEY:
    st.error("🚨 API key is missing. Please set PAGESPEED_API_KEY in .env.")
    st.stop()

API_ENDPOINT = "https://www.googleapis.com/pagespeedonline/v5/runPagespeed"

# Fetch data from PageSpeed API
def fetch_pagespeed_data(url):
    params = {"url": url, "key": API_KEY, "strategy": "mobile"}
    response = requests.get(API_ENDPOINT, params=params)
    return response.json()

# Extract relevant audit summary
def extract_summary(data):
    try:
        title = data["lighthouseResult"]["finalUrl"]
        seo_score = data["lighthouseResult"]["categories"]["seo"]["score"] * 100
        perf_score = data["lighthouseResult"]["categories"]["performance"]["score"] * 100
        return {
            "title": title,
            "seo_score": int(seo_score),
            "perf_score": int(perf_score),
        }
    except KeyError:
        return None

# Generate Word report
def generate_docx(summary_data):
    doc = Document()
    doc.add_heading("SEO Audit Report", 0)
    doc.add_paragraph(f"URL: {summary_data['title']}")
    doc.add_paragraph(f"SEO Score: {summary_data['seo_score']}%")
    doc.add_paragraph(f"Performance Score: {summary_data['perf_score']}%")
    
    output = io.BytesIO()
    doc.save(output)
    output.seek(0)
    return output

# ───── Streamlit UI ─────
st.set_page_config(page_title="Mini SEO Auditor", page_icon="🔍")
st.title("🔍 Mini SEO Auditor")
st.markdown("Check SEO and Performance scores using Google PageSpeed API.")

site_url = st.text_input("Enter website URL to audit (include http/https):", "")

if site_url:
    if st.button("Run Audit"):
        with st.spinner("Fetching data from Google PageSpeed..."):
            full_json = fetch_pagespeed_data(site_url)
            summary_data = extract_summary(full_json)

            if summary_data:
                st.success("Audit completed successfully!")
                st.subheader("📈 SEO Summary")
                st.write(summary_data)

                # Word Download
                docx_file = generate_docx(summary_data)
                st.download_button(
                    "📄 Download Word Report",
                    docx_file,
                    file_name="seo_report.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                )

                # JSON Download
                json_file = io.BytesIO(json.dumps(full_json, indent=2).encode("utf-8"))
                st.download_button(
                    "📁 Download Full JSON Report",
                    json_file,
                    file_name="seo_report.json",
                    mime="application/json"
                )
            else:
                st.error("Failed to extract SEO data from the API response.")
